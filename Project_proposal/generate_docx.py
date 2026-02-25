"""Generate Project Proposal docx from the markdown content."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Helper functions ──

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, size=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = size
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing after table
    return table

def add_bold_then_normal(bold_text, normal_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run_b = p.add_run(bold_text)
    run_b.bold = True
    run_b.font.size = Pt(11)
    run_n = p.add_run(normal_text)
    run_n.font.size = Pt(11)
    return p

# ═══════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════
title = doc.add_heading('EE3070 Project Proposal', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

add_para('Group No:  L ____  Gp ____', bold=True, size=Pt(12))
add_para('AI-Bot — Voice-Controlled Desktop AI Assistant with Computer Automation',
         bold=True, size=Pt(13), space_after=Pt(12))

# Member table
make_table(
    ['Member Name', 'Group Leader (√)', 'Student ID', 'Signature'],
    [['', '', '', ''],
     ['', '', '', ''],
     ['', '', '', ''],
     ['', '', '', '']],
    col_widths=[5, 3.5, 3.5, 4]
)

# ═══════════════════════════════════════════════════════════
# SECTION 1
# ═══════════════════════════════════════════════════════════
add_heading_text('Section 1: Description of Targeted Scenario and Project Goals', level=1)

add_heading_text('Background', level=2)
add_para(
    'With the rapid development of artificial intelligence, voice assistants such as Amazon Alexa, '
    'Google Home, and Apple Siri have become common household devices. However, these commercial '
    'products are limited to cloud-based services and lack the ability to directly control a user\'s '
    'personal computer. Meanwhile, desktop users frequently perform repetitive tasks — opening '
    'applications, managing files, adjusting system settings — that could be automated through '
    'natural voice commands.'
)
add_para(
    'There is a gap in the market for a compact desktop AI assistant that combines hardware-based '
    'voice interaction with direct computer control capabilities. Existing solutions either require '
    'cloud processing (raising privacy concerns) or are purely software-based (lacking a dedicated '
    'physical interface with display and tactile feedback).'
)

add_heading_text('Identified Problems', level=2)
add_numbered('Lack of physical AI interface for desktop control: Current voice assistants cannot directly manage files, launch applications, or control system settings on a personal computer.')
add_numbered('Privacy concerns with cloud-only voice assistants: Commercial voice assistants send all audio to remote servers, posing risks for sensitive work environments.')
add_numbered('Inefficient desktop workflows: Users repeatedly perform manual operations (opening apps, adjusting volume, organizing files) that could be streamlined through voice automation.')
add_numbered('No affordable, customizable desktop AI hardware: Existing smart speakers are closed ecosystems with no customization or local AI processing options.')

add_heading_text('Ethical Constraints', level=2)
add_numbered('Privacy and data security: The system transmits user audio to cloud-based APIs (Whisper for speech recognition, LLM for dialogue, Edge-TTS for synthesis) for processing. To mitigate privacy risks, audio is transmitted over encrypted HTTPS connections, is not stored persistently on the server side, and is used solely for real-time processing. Users are clearly informed that voice data is uploaded to the cloud, and the device only begins recording after an explicit wake-up trigger, preventing unintended audio capture.')
add_numbered('User consent and control: All computer control operations (file deletion, application management) require explicit user confirmation through voice or physical touch interaction (tap-to-confirm). A dedicated "chat-only mode" disables all computer control for safety.')
add_numbered('Transparency: The system clearly indicates its operational state (recording, processing, executing) through the LCD display and LED indicators, ensuring users always know what the device is doing.')
add_numbered('Responsible AI use: The AI dialogue engine operates within defined boundaries and does not perform destructive or unauthorized system operations. Sensitive operations are sandboxed.')

add_heading_text('Impact of the Solution', level=2)
add_bullet('Accessibility: Provides a hands-free computer control interface that benefits users with mobility impairments or those who prefer voice interaction for multitasking.')
add_bullet('Productivity: Automates repetitive desktop tasks, enabling users to manage files, launch applications, and control system settings through simple voice commands, improving workflow efficiency.')
add_bullet('Educational value: The project integrates multiple engineering disciplines — embedded systems, PCB design, wireless communication, AI, and software development — serving as a comprehensive learning experience and demonstrating how IoT hardware and cloud AI services can be combined into a practical consumer product.')
add_bullet('Affordable and customizable alternative: Compared to commercial smart speakers that are locked into proprietary ecosystems, this project offers a cost-effective, fully customizable desktop AI assistant tailored to individual user needs, with the flexibility to choose different AI models and extend functionality.')

add_heading_text('Project Goals', level=2)
add_numbered('Design and fabricate a custom PCB integrating ESP32-S3 microcontroller with audio input/output, display, motion sensing, and power management.')
add_numbered('Implement real-time voice interaction: user speaks to the device, speech is recognized, AI generates a response, and the response is played back through the speaker.')
add_numbered('Enable voice-controlled computer automation: file management, application control, and system settings adjustment through natural language commands.')
add_numbered('Provide a standby desktop widget mode displaying time, weather, and status information on the built-in LCD screen.')
add_numbered('Implement physical interaction features including tap-to-confirm (capacitive touch) and shake-to-trigger (accelerometer-based) functions.')

# ═══════════════════════════════════════════════════════════
# SECTION 2
# ═══════════════════════════════════════════════════════════
add_heading_text('Section 2: Description of System Functions', level=1)

add_heading_text('System Architecture Overview', level=2)
add_para(
    'The system consists of two main parts connected via WiFi (WebSocket):\n'
    '(1) Hardware Device: ESP32-S3 + INMP441 Mic + MAX98357A Amp + ST7789 LCD + MPU6050 IMU '
    '+ Capacitive Touch + WS2812B LEDs + TP4056 Charger + AMS1117 LDO + Li-ion Battery\n'
    '(2) PC Server: FastAPI Backend with Whisper (STT), Edge-TTS (TTS), LLM Engine (GPT-4/Claude), '
    'and NanoBot (Lightweight Computer Control Agent)',
    size=Pt(10)
)

# -- Function 1 --
add_heading_text('Function 1: Voice Interaction (MVP)', level=2)
add_bold_then_normal('Modules used: ',
    'INMP441 (I2S digital microphone), MAX98357A (I2S Class-D amplifier), 3W speaker, '
    'ESP32-S3 (WiFi), PC server (Whisper + Edge-TTS + LLM)')
add_bold_then_normal('How it works:', '')
add_numbered('User speaks to the device; INMP441 captures audio via I2S interface.')
add_numbered('ESP32-S3 streams audio data to the PC server over WiFi (WebSocket).')
add_numbered('Server runs Whisper speech recognition to convert audio to text.')
add_numbered('Text is sent to the LLM (GPT-4/Claude) for intent understanding and response generation.')
add_numbered('Response text is converted to audio using Edge-TTS.')
add_numbered('Audio is streamed back to ESP32-S3, which plays it through MAX98357A and speaker.')
add_bold_then_normal('User interaction: ',
    'User simply speaks to the device naturally. The device shows recording/processing/playback status on the LCD screen.')

# -- Function 2 --
add_heading_text('Function 2: Computer Control via Voice (MVP)', level=2)
add_bold_then_normal('Modules used: ',
    'PC server (NanoBot lightweight computer control agent), LLM intent recognition')
add_bold_then_normal('How it works:', '')
add_numbered('After speech recognition, the LLM identifies the user\'s intent (e.g., "open VS Code", "create a folder on Desktop").')
add_numbered('If the intent is a computer control command, the server routes it to NanoBot.')
add_numbered('NanoBot is a lightweight, easily customizable agent that executes computer control tasks including file operations, application management, and system commands.')
add_numbered('Execution result is returned to the user via voice feedback.')
add_bold_then_normal('Why NanoBot: ',
    'NanoBot was chosen for its lightweight architecture and ease of modification. OpenClaw\'s large codebase (700+ skills) proved difficult to customize for our specific project needs, while NanoBot provides a compact, developer-friendly framework that is straightforward to extend and adapt.')
add_bold_then_normal('Supported operations:', '')
add_bullet('File management: create, delete, move, rename, search files/folders')
add_bullet('Application control: open, close, switch applications')
add_bullet('System operations: adjust volume, brightness, take screenshots, lock screen')
add_bold_then_normal('Safety: ',
    'Destructive operations require user confirmation via tap (capacitive touch) or voice ("confirm" / "cancel").')

# -- Function 3 --
add_heading_text('Function 3: Desktop Widget Mode (MVP)', level=2)
add_bold_then_normal('Modules used: ',
    'ST7789 1.69" IPS LCD (SPI interface), ESP32-S3 (NTP time sync, weather API)')
add_bold_then_normal('How it works:', '')
add_numbered('In standby mode, the LCD displays current time (synced via NTP), weather information (from weather API), and upcoming reminders.')
add_numbered('The display updates periodically with minimal power consumption.')
add_numbered('UI is rendered using the LVGL embedded GUI framework.')
add_bold_then_normal('User interaction: ',
    'The device serves as a desktop clock/weather station when not actively in conversation.')

# -- Function 4 --
add_heading_text('Function 4: Physical Interaction', level=2)
add_bold_then_normal('Modules used: ',
    'ESP32-S3 built-in capacitive touch sensor (IO7), MPU6050 6-axis IMU (I2C), WS2812B RGB LEDs')
add_bullet('Tap interaction: Capacitive touch pad on the top sub-board detects single tap (confirm), double tap (reject), and triple tap (interrupt). Replaces mechanical buttons for a cleaner design.')
add_bullet('Shake to trigger: MPU6050 accelerometer detects shaking motion, triggering fun features (daily fortune, random quotes, decision maker).')
add_bullet('LED feedback: WS2812B RGB LEDs provide ambient lighting and status indication (breathing effect during standby, color changes during recording/processing).')

# -- Function 5 --
add_heading_text('Function 5: Power Management', level=2)
add_bold_then_normal('Modules used: ',
    'TP4056 (Li-ion charge controller), AMS1117-3.3 (LDO voltage regulator), USB Type-C (16-pin), 3.7V Li-ion battery (1000–2000mAh)')
add_numbered('USB Type-C provides 5V input for charging via TP4056 (1A charge current, set by 1.2KΩ PROG resistor).')
add_numbered('TP4056 manages battery charging with LED status indicators (red = charging, green = full).')
add_numbered('AMS1117-3.3 regulates battery voltage to stable 3.3V for all digital components.')
add_numbered('The device operates cordlessly on battery power for portable use.')

# -- BOM --
add_heading_text('Bill of Materials (BOM)', level=2)
add_para('The following component list is derived from the actual schematic design (v2.0).', italic=True, size=Pt(10))

add_para('Main Board — Active Components (ICs & Modules)', bold=True, size=Pt(11))
make_table(
    ['Ref', 'Component', 'Part Number / Value', 'Package', 'Qty', 'Function'],
    [
        ['U1', 'MCU Module', 'ESP32-S3-WROOM-1-N16R8', 'Module', '1', 'Main controller (WiFi/BT, 16MB Flash + 8MB PSRAM)'],
        ['U2', 'MEMS Microphone', 'INMP441 (EV_INMP441)', 'Module', '1', 'I2S digital audio capture'],
        ['U3', 'Audio Amplifier', 'MAX98357AETE+T', 'QFN-16', '1', 'I2S Class-D amplifier, drives speaker'],
        ['U4', 'Charge Controller', 'TP4056', 'SOP-8', '1', 'Li-ion battery charge management (1A)'],
        ['U5', 'Voltage Regulator', 'AMS1117-3.3', 'SOT-223', '1', '3.3V LDO (battery to 3.3V)'],
        ['U7', '6-Axis IMU', 'ZY-MPU-6050', 'Stamp-hole', '1', 'Accelerometer + gyroscope'],
    ],
    col_widths=[1.2, 2.5, 3.5, 2, 1, 5.5]
)

add_para('Main Board — Connectors', bold=True, size=Pt(11))
make_table(
    ['Ref', 'Component', 'Part Number', 'Qty', 'Function'],
    [
        ['J1', 'Battery Connector', 'B2B-PH-K-S (JST PH 2-pin)', '1', 'Li-ion battery connection'],
        ['J2', 'USB Type-C', 'KH-TYPE-C-16P-N (16-pin)', '1', '5V charging and USB data/debug'],
        ['J3', 'Display FPC', 'AFC01-S12FCA-00 (12-pin)', '1', 'ST7789 LCD connection'],
        ['P1', 'Speaker Connector', 'PZ254V-11-02P (2-pin)', '1', '3W 4Ω speaker connection'],
        ['H1', 'Expansion Header', '2.54mm 1×5P Female', '1', 'IO1, IO2, TXD0, RXD0, IO39'],
        ['H2', 'Expansion Header', '2.54mm 1×5P Female', '1', 'IO13, IO45, IO48, IO47, IO21'],
        ['H4', 'Expansion Header', '2.54mm 1×5P Female', '1', 'IO3, 3V3, GND (sub-board link)'],
    ],
    col_widths=[1.2, 2.5, 4, 1, 5]
)

add_para('Main Board — Passive Components', bold=True, size=Pt(11))
make_table(
    ['Ref', 'Component', 'Value', 'Pkg', 'Qty', 'Function'],
    [
        ['C1, C2', 'Capacitor', '10µF', '0805', '2', 'ESP32-S3 power decoupling'],
        ['C3', 'Capacitor', '100nF', '0603', '1', 'ESP32-S3 power decoupling'],
        ['C5', 'Capacitor', '100nF', '0603', '1', 'INMP441 VDD decoupling'],
        ['C7', 'Capacitor', '100nF', '0603', '1', 'MAX98357A VBAT decoupling'],
        ['C8', 'Capacitor', '100nF', '0603', '1', 'ST7789 backlight / power decoupling'],
        ['C9–C11', 'Capacitor', 'Various', '0603/0805', '3', 'TP4056 & AMS1117 filtering'],
        ['C13', 'Capacitor', '100nF', '0603', '1', 'MPU6050 VCC decoupling'],
        ['R1, R2', 'Resistor', '5.1kΩ', '0603', '2', 'USB CC1/CC2 pull-down'],
        ['R3, R4', 'Resistor', 'Limiting', '0603', '2', 'Charge LED current limiting'],
        ['R6', 'Resistor', '4.7Ω', '0603', '1', 'ST7789 backlight LED limiting'],
        ['R12', 'Resistor', '100kΩ', '0603', '1', 'INMP441 L/R channel select'],
        ['R_PROG', 'Resistor', '1.2kΩ', '0603', '1', 'TP4056 charge current (1A)'],
        ['D1', 'LED', 'Red', '0603', '1', 'Charging indicator (CHRG)'],
        ['D2', 'LED', 'Green', '0603', '1', 'Standby/full indicator (STDBY)'],
    ],
    col_widths=[1.5, 1.8, 1.5, 1.5, 1, 5]
)

add_para('Main Board — Switches', bold=True, size=Pt(11))
make_table(
    ['Ref', 'Component', 'Part Number', 'Qty', 'Function'],
    [
        ['SW1', 'Slide Switch', 'SK12D07VG4', '1', 'Main power on/off'],
        ['—', 'Reset Button', 'Tactile switch', '1', 'ESP32-S3 EN reset (RC delay)'],
    ],
    col_widths=[1.2, 2.5, 3, 1, 5]
)

add_para('Main Board — Other External Components', bold=True, size=Pt(11))
make_table(
    ['Component', 'Qty', 'Function'],
    [
        ['3W 4Ω Speaker', '1', 'Audio output'],
        ['1.69" IPS TFT LCD (ST7789)', '1', 'Display (via FPC to J3)'],
        ['3.7V Li-ion Battery (1000–2000mAh)', '1', 'Portable power source'],
        ['WS2812B LED Strip', '1', 'Ambient lighting / status (IO38)'],
    ],
    col_widths=[5, 1.5, 7]
)

add_para('Touch Sub-Board', bold=True, size=Pt(11))
make_table(
    ['Ref', 'Component', 'Value / Part', 'Qty', 'Function'],
    [
        ['R5', 'Resistor', '4.7kΩ', '1', 'ESD protection for touch pad'],
        ['TP1–TP4', 'Test Points', '—', '4', 'TOUCH_MAIN, GND, 3V3, LED_DIN'],
        ['TP5–TP6', 'Test Points', '—', '2', 'Touch copper pad connections'],
        ['—', 'Copper Pad', 'PCB copper pour', '1', 'Capacitive touch area (8–12mm)'],
    ],
    col_widths=[1.5, 2, 2.5, 1, 5.5]
)

# ═══════════════════════════════════════════════════════════
# SECTION 3
# ═══════════════════════════════════════════════════════════
add_heading_text('Section 3: Project Planning Schedule', level=1)

make_table(
    ['Week', 'Tasks to be completed'],
    [
        ['6', 'Submit project proposal. Receive PCB boards and components. Begin soldering main board and touch sub-board.'],
        ['7', 'Complete PCB soldering. Power-on test (verify 3.3V supply). Test ESP32-S3 basic functionality (WiFi, serial debug). Begin individual module testing (microphone, speaker, display).'],
        ['8', 'Complete hardware module testing: INMP441 audio capture, MAX98357A audio playback, ST7789 display output. Test MPU6050 motion detection and capacitive touch sensing. Begin firmware development (WiFi communication, I2S audio driver).'],
        ['9', 'Integrate all hardware modules for full system test. Set up PC server environment: install Python, FastAPI, Whisper, Edge-TTS. Implement WebSocket communication between ESP32-S3 and PC server. Begin end-to-end voice pipeline.'],
        ['10', 'Complete voice interaction pipeline: speech recognition (Whisper) + AI response (LLM) + text-to-speech (Edge-TTS) + playback. Integrate NanoBot for computer control. Implement basic computer control commands.'],
        ['11', 'Implement standby widget mode (time/weather display). Implement physical interaction (tap-to-confirm, shake-to-trigger). Implement LED ambient lighting effects. System stability testing and bug fixing.'],
        ['12', 'Complete 3D-printed enclosure assembly. Full system integration testing. Prepare demonstration: voice dialogue, computer control, widget mode, physical interaction. Write final report.'],
        ['13', 'Final presentation and demonstration.'],
    ],
    col_widths=[1.5, 14]
)

# ═══════════════════════════════════════════════════════════
# SECTION 4
# ═══════════════════════════════════════════════════════════
add_heading_text('Section 4: Division of Labor', level=1)
add_para('(Please fill in member names and Student IDs)', italic=True)

make_table(
    ['Member', 'Responsibilities'],
    [
        ['Member 1\n(Name, SID)', 'Hardware design & PCB: Schematic design in LCEDA (all 7 modules), PCB layout and routing, soldering, hardware debugging. 3D enclosure design (Autodesk Fusion).'],
        ['Member 2\n(Name, SID)', 'Firmware development: ESP32-S3 firmware (Arduino/MicroPython) — WiFi communication, I2S audio driver, SPI display driver, touch/IMU sensor integration, LED control, device state machine.'],
        ['Member 3\n(Name, SID)', 'PC server backend: Python + FastAPI server, WebSocket communication handler, Whisper speech recognition integration, Edge-TTS voice synthesis, LLM API integration (GPT-4/Claude), NanoBot computer control integration.'],
        ['Member 4\n(Name, SID)', 'System integration & testing: End-to-end testing of voice pipeline, computer control validation, standby display features, physical interaction testing, demo preparation, documentation and report writing.'],
    ],
    col_widths=[3, 13]
)

# ═══════════════════════════════════════════════════════════
# SECTION 5 — Gantt Chart as table
# ═══════════════════════════════════════════════════════════
add_heading_text('Section 5: Gantt Chart', level=1)

gantt_headers = ['Task / Member', 'Wk6', 'Wk7', 'Wk8', 'Wk9', 'Wk10', 'Wk11', 'Wk12', 'Wk13']
gantt_rows = [
    ['PCB Soldering & Power Test (M1)',         '██', '█',  '',   '',   '',   '',   '',   ''],
    ['Module Hardware Testing (M1 & M2)',       '',   '██', '██', '',   '',   '',   '',   ''],
    ['Firmware Dev: WiFi, I2S, SPI, Touch (M2)','',   '',   '█',  '██', '█',  '',   '',   ''],
    ['PC Server: FastAPI, Whisper, TTS, LLM (M3)','', '',   '',   '██', '██', '',   '',   ''],
    ['NanoBot Integration (M3)',                '',   '',   '',   '',   '██', '█',  '',   ''],
    ['Widget Mode & Physical Interaction (M2&4)','',  '',   '',   '',   '',   '██', '█',  ''],
    ['3D Enclosure Design & Assembly (M1)',     '',   '',   '',   '',   '█',  '█',  '█',  ''],
    ['System Integration & E2E Testing (M4)',   '',   '',   '',   '',   '',   '█',  '██', ''],
    ['Report Writing & Demo Prep (All)',        '',   '',   '',   '',   '',   '',   '██', ''],
    ['Final Presentation (All)',                '',   '',   '',   '',   '',   '',   '',   '██'],
]

table = doc.add_table(rows=1 + len(gantt_rows), cols=len(gantt_headers))
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for i, h in enumerate(gantt_headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data
for r_idx, row in enumerate(gantt_rows):
    for c_idx, val in enumerate(row):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if c_idx == 0:
            run.font.size = Pt(8)
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if val:
                # Shade filled cells
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): '4472C4'
                })
                shading.append(shading_elem)
                run.font.color.rgb = RGBColor(255, 255, 255)

# Set column widths
gantt_widths = [5.5, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2, 1.2]
for i, w in enumerate(gantt_widths):
    for row in table.rows:
        row.cells[i].width = Cm(w)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# SECTION 6
# ═══════════════════════════════════════════════════════════
add_heading_text('Section 6: Summarize the Current Progress', level=1)

add_heading_text('Overall Status', level=2)
add_para(
    'The project is currently in the hardware fabrication stage. Schematic design and PCB layout '
    'have been completed and verified. PCB boards have been ordered for fabrication and all '
    'components have been purchased. 3D enclosure design is in progress.'
)

add_heading_text('Completed Work', level=2)

add_para('1. Project Planning and Research (Week 1–2)', bold=True)
add_bullet('Established project goals and defined core functions (voice interaction, computer control, desktop widget, physical interaction).')
add_bullet('Researched and selected key software frameworks: NanoBot (lightweight computer control agent), Whisper (speech recognition), Edge-TTS (speech synthesis).')
add_bullet('Compared NeoAI, OpenClaw, and NanoBot for computer control. Selected NanoBot for its lightweight architecture and ease of customization.')
add_bullet('Created detailed function implementation documents covering all 9 planned system functions.')

add_para('2. Component Selection (Week 2)', bold=True)
add_bullet('Selected ESP32-S3-WROOM-1-N16R8 as the main MCU (dual-core, WiFi/Bluetooth, 16MB Flash + 8MB PSRAM).')
add_bullet('Selected INMP441 (I2S digital microphone) and MAX98357A (I2S Class-D amplifier) for audio.')
add_bullet('Selected 1.69" IPS TFT with ST7789 driver (SPI interface) for display.')
add_bullet('Selected MPU6050 6-axis IMU for motion detection.')
add_bullet('Selected TP4056 + AMS1117-3.3 for power management.')
add_bullet('Designed capacitive touch interaction using ESP32-S3\'s built-in touch sensor (replacing external TTP223 chip, reducing BOM by 5–7 components).')

add_para('3. Pin Assignment Planning (Week 2)', bold=True)
add_bullet('Completed full GPIO pin assignment for ESP32-S3, accounting for N16R8 module restrictions (GPIO22–37 occupied by internal Octal SPI Flash/PSRAM).')
add_bullet('Documented all pin assignments: I2S audio (IO14–18, IO8), SPI display (IO9–12, IO46), I2C IMU (IO5–6), touch (IO7), LED (IO38), USB (IO19–20).')

add_para('4. Schematic Design in LCEDA (Week 2–5)', bold=True)
add_bullet('Completed schematics for all 7 hardware modules:')
add_bullet('ESP32-S3 minimum system (power decoupling, EN reset, IO0 boot, USB Type-C)', level=1)
add_bullet('INMP441 microphone (I2S interface)', level=1)
add_bullet('MAX98357A amplifier (I2S interface, gain configuration)', level=1)
add_bullet('ST7789 display (SPI interface, backlight driver)', level=1)
add_bullet('Power management (TP4056 charge IC + AMS1117-3.3 LDO)', level=1)
add_bullet('MPU6050 gyroscope (I2C interface, stamp-hole module)', level=1)
add_bullet('Capacitive touch system (ESP32-S3 built-in, with ESD protection resistors)', level=1)
add_bullet('Passed ERC (Electrical Rules Check) with no errors.')

add_para('5. PCB Design (Week 5–6)', bold=True)
add_bullet('Completed PCB layout and routing for the main board.')
add_bullet('Designed a separate touch sub-board (top panel) with capacitive touch pad and ESD protection.')
add_bullet('Sub-board simplified to a pure touch sensing board (copper pad + ESD resistor).')
add_bullet('Passed DRC (Design Rules Check) for both main board and sub-board.')

add_para('6. Procurement (Week 6)', bold=True)
add_bullet('Ordered PCB fabrication through JLCPCB (main board + touch sub-board).')
add_bullet('Purchased all electronic components.')

add_para('7. 3D Enclosure Design (In Progress)', bold=True)
add_bullet('Started designing a 3D-printed enclosure using Autodesk Fusion.')

add_heading_text('Documentation Produced', level=2)
add_bullet('README.md — Project overview and feature descriptions')
add_bullet('CLAUDE.md — Project guidelines and pin assignments')
add_bullet('CHANGELOG.md — Detailed daily work log (9 days of progress)')
add_bullet('功能讨论区/功能实现讨论.md — Detailed implementation plan for all 9 functions')
add_bullet('功能讨论区/task.md — MVP task checklist with 8 development phases')
add_bullet('功能讨论区/openClaw.md — Computer control framework research (OpenClaw, NeoAI, NanoBot comparison)')
add_bullet('原理图设计/01–07, 09 — Step-by-step schematic design tutorials for each module')
add_bullet('Component datasheets collected in 元件资料区/ and 元件TXT/')

# ── Save ──
output_path = '/Users/mandy/Documents/GitHub/AI-bot/Project_proposal/Project_Proposal.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
